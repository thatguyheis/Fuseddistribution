from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("deliverables/nicholas-hughes-general-manager-resume-fused.docx")

NAVY = RGBColor(24, 55, 83)
INK = RGBColor(32, 38, 45)
GRAY = RGBColor(84, 92, 101)
LIGHT = RGBColor(205, 213, 220)
WHITE = RGBColor(255, 255, 255)
FONT = "Aptos"


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "nil")


def set_run(run, size=10.2, bold=False, color=INK, italic=False, all_caps=False):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.all_caps = all_caps


def add_paragraph(doc, text="", *, size=10.2, bold=False, color=INK, italic=False,
                  before=0, after=3, line=1.04, align=WD_ALIGN_PARAGRAPH.LEFT,
                  keep_next=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep_next
    if text:
        set_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_rule(paragraph, color="183753", size="10", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_section_heading(doc, text):
    p = add_paragraph(doc, text.upper(), size=10.3, bold=True, color=NAVY,
                      before=7, after=4, line=1.0, keep_next=True)
    add_rule(p, color="B8C5CF", size="5", space="3")
    return p


def add_bullet(doc, text, *, after=2.2):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.20)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.02
    p.paragraph_format.keep_together = True
    set_run(p.add_run(text), size=9.75, color=INK)
    return p


def add_role(doc, company, location, title, dates, bullets):
    p = add_paragraph(doc, before=5, after=0.5, line=1.0, keep_next=True)
    set_run(p.add_run(company), size=10.3, bold=True, color=NAVY)
    if location:
        set_run(p.add_run(f" | {location}"), size=9.7, color=GRAY)
    p2 = add_paragraph(doc, after=2.2, line=1.0, keep_next=True)
    set_run(p2.add_run(title), size=9.9, bold=True, color=INK)
    set_run(p2.add_run(f"  |  {dates}"), size=9.6, color=GRAY)
    for bullet in bullets:
        add_bullet(doc, bullet)


def set_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    set_run(run, size=8.5, color=GRAY)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Named resume_compact override: tighter margins preserve a readable two-page resume.
    section.top_margin = Inches(0.57)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.04

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet.font.size = Pt(9.75)
    bullet.paragraph_format.left_indent = Inches(0.20)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.space_after = Pt(2.2)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(p.add_run("Nicholas Hughes  |  "), size=8.5, color=GRAY)
    set_page_field(p)


def build_resume():
    doc = Document()
    configure_document(doc)

    name = add_paragraph(doc, "NICHOLAS HUGHES", size=22, bold=True, color=NAVY,
                         after=1, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    title = add_paragraph(doc, "OPERATIONS & PEOPLE LEADER | GENERAL MANAGER CANDIDATE",
                          size=10.3, bold=True, color=GRAY, after=2.5, line=1.0,
                          align=WD_ALIGN_PARAGRAPH.CENTER)
    contact = add_paragraph(doc, "(360) 213-6035  |  n_hughes@apple.com  |  Portland, Oregon Area",
                            size=9.5, color=INK, after=1.5, line=1.0,
                            align=WD_ALIGN_PARAGRAPH.CENTER)
    website = add_paragraph(doc, "fuseddistribution.com", size=9.3, bold=True, color=NAVY,
                            after=5, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_rule(website, color="183753", size="9", space="4")

    add_section_heading(doc, "Executive Profile")
    add_paragraph(
        doc,
        "Operations leader and entrepreneur with 13+ years of experience turning ideas into practical systems, "
        "coaching teams, improving performance, and coordinating complex work across functions. Built and operates "
        "Fused Distribution and Fused Reserve, with hands-on responsibility for digital operations, CRM and lead "
        "workflows, AI automation, inventory planning, sourcing, market analysis, and partner development. Brings "
        "leadership experience across teams of 20+ employees, coaching programs serving 50+ people, vendor programs, "
        "KPI reporting, budget tracking, and customer experience.",
        size=10.0, after=4.5, line=1.05,
    )

    add_section_heading(doc, "Leadership Impact")
    impacts = [
        "Built and led a repeatable peer coaching program that improved participant performance scores by more than 20%.",
        "Led 20+ employees across three departments and increased store sales 5% in the first month and 10% by month three.",
        "Coached a 50+ person team on updated procedures, reducing filing errors and improving operational throughput.",
        "Built Fused Distribution and Fused Reserve operating systems spanning website delivery, CRM, lead generation, AI content automation, inventory, sourcing, and market analysis.",
    ]
    for impact in impacts:
        add_bullet(doc, impact, after=2.4)

    add_section_heading(doc, "Core Capabilities")
    capabilities = [
        "Operational Strategy & Execution  |  People Leadership, Coaching & Accountability",
        "CRM, Lead Generation & Customer Journeys  |  KPI Dashboards & Performance Reviews",
        "Inventory, Sourcing & Market Analysis  |  Budget Tracking & Forecasting Support",
        "AI Automation & Workflow Design  |  Partnerships & Vendor Operations",
    ]
    for value in capabilities:
        add_paragraph(doc, value, size=9.55, bold=True, after=1.8, line=1.0)

    add_section_heading(doc, "Professional Experience")
    add_role(
        doc, "Fused Distribution / Fused Reserve", "Portland, Oregon", "Founder & Operations Lead", "Apr 2026 - Present",
        [
            "Built and operates Fused Distribution, translating business ideas into service offerings, pricing, website experiences, SOPs, production workflows, and customer acquisition systems.",
            "Manages CRM-style lead operations including website lead capture, inquiry routing, pipeline organization, newsletter segmentation, follow-up workflows, and customer journey improvements.",
            "Designed AI-assisted production systems that research, create, validate, and publish community-focused content, then generate narrated vertical videos and reels with captions, licensed media, and quality controls.",
            "Built Fused Reserve's operating model and manages physical silver inventory, sourcing channels, purchase qualification, demand forecasting, fulfillment planning, and pricing decisions tied to spot and product premiums.",
            "Develops dealer and business partnerships while analyzing market pricing, customer segments, product demand, sourcing performance, margins, and inventory levels to improve reach and guide growth decisions.",
        ],
    )
    add_role(
        doc, "Apple", "", "Project Coordinator, Customer Relations Back Office", "Jan 2026 - Present",
        [
            "Coordinate vendor operations for AI evaluation programs from onboarding through final delivery, aligning external linguists, small vendors, and internal program managers across multiple concurrent workstreams.",
            "Own structured performance trackers that serve as the source of truth for vendor capacity, availability, task progress, spend, risk, and overall program health.",
            "Translate leadership priorities and technical findings into project plans and decision-ready reporting; research vendor capabilities, identify operational gaps, and recommend sourcing actions.",
            "Track vendor spend against program targets, support forecasting, escalate delivery risks, and use AI tools to improve research and data workflows.",
        ],
    )

    doc.add_page_break()

    add_section_heading(doc, "Professional Experience, Continued")
    add_role(
        doc, "Apple", "", "Senior Advisor, Mac+ Plus / MSS / iOS Support", "Mar 2023 - Jan 2026",
        [
            "Served as the final escalation point for complex customer and vendor issues, balancing service quality, policy, speed, and customer trust in a high volume environment.",
            "Analyzed performance data and recurring operational patterns, then partnered with Engineering, Dispatch, Sales, and leadership to implement corrective action and durable process improvements.",
            "Facilitated team meetings and performance initiatives that improved team metrics while managing a demanding caseload and shifting priorities.",
        ],
    )

    add_role(
        doc, "Apple", "", "Social Media Response Advisor", "Aug 2016 - Mar 2023",
        [
            "Built and led a peer coaching program from concept through execution, producing performance score improvements of more than 20% for participating advisors.",
            "Created solutions for complex customer issues with no established resolution path and distributed them across the team to improve consistency, quality, and decision speed.",
            "Coordinated across teams during holiday surges and major product launches; also served as lead for the Battery Replacement team when operational needs shifted.",
        ],
    )
    add_role(
        doc, "Apple", "", "Mac and iPhone Technical Advisor", "Aug 2015 - Aug 2016",
        [
            "Led the team in customer satisfaction, repair quality, and escalation results.",
            "Developed meeting content and structured practice sessions that reduced average call handle time across the group.",
        ],
    )
    add_role(
        doc, "Microsoft", "Portland, Oregon", "Microsoft Expert (Store Manager)", "May 2015 - Aug 2015",
        [
            "Led more than 20 employees across three departments, setting expectations, coaching performance, and coordinating daily execution in a high visibility retail environment.",
            "Increased store sales 5% in the first month and 10% by month three through disciplined performance management and customer focused execution.",
        ],
    )
    add_role(
        doc, "Best Buy", "Portland, Oregon", "Senior Sales Consultant (Supervisor)", "Jul 2014 - May 2015",
        [
            "Led weekly team training on product knowledge, sales execution, and merchandising standards.",
            "Built employee schedules and clarified operating roles to improve floor coverage, efficiency, and team accountability.",
        ],
    )
    add_role(
        doc, "Geek Squad", "Portland, Oregon", "Operations Agent (Back Office Lead)", "Mar 2013 - Jul 2014",
        [
            "Managed back office operations including inbound coordination, appointment scheduling, workflow prioritization, and operational reporting.",
            "Led weekly coaching for more than 50 employees on procedures and quality standards, reducing filing errors and improving throughput.",
        ],
    )

    add_section_heading(doc, "Education & Professional Development")
    education = [
        ("Google Project Management Professional Certificate", "Google / Coursera"),
        ("Google AI Certificate", "Google / Coursera"),
        ("Associate of Arts Transfer", "Clark College, Vancouver, Washington"),
    ]
    for credential, issuer in education:
        p = add_paragraph(doc, after=2, line=1.0)
        set_run(p.add_run(credential), size=9.8, bold=True, color=INK)
        set_run(p.add_run(f"  |  {issuer}"), size=9.5, color=GRAY)

    add_section_heading(doc, "Technology")
    add_paragraph(
        doc,
        "CRM and lead workflows | Reporting dashboards and structured trackers | Cloudflare and web operations | "
        "Remotion and Buffer | Microsoft Office and Google Workspace | Claude, ChatGPT, Codex, and Gemini",
        size=9.7, after=0, line=1.0,
    )

    doc.core_properties.title = "Nicholas Hughes - General Manager Resume with Fused Distribution Experience"
    doc.core_properties.subject = "Targeted resume for General Manager, Simple Lawns & Landscape Design"
    doc.core_properties.author = "Nicholas Hughes"
    doc.core_properties.keywords = "general manager, operations, leadership, KPIs, coaching, budgeting, process improvement"
    doc.core_properties.comments = ""

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build_resume()
